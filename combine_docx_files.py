import os
from tkinter import filedialog, Tk
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def select_files():
    root = Tk()
    root.withdraw()  # Hide the root window
    file_paths = filedialog.askopenfilenames(
        title="Select .docx files to combine",
        filetypes=[("Word Documents", "*.docx")])
    return list(file_paths)

def append_paragraphs(target_document, source_document):
    for paragraph in source_document.paragraphs:
        target_paragraph = target_document.add_paragraph()
        for run in paragraph.runs:
            target_run = target_paragraph.add_run(run.text)
            target_run.bold = run.bold
            target_run.italic = run.italic
            target_run.underline = run.underline
            target_run.font.name = run.font.name
            if run.font.size:
                target_run.font.size = run.font.size
            target_run.font.color.rgb = run.font.color.rgb

def append_tables(target_document, source_document):
    for table in source_document.tables:
        target_table = target_document.add_table(rows=0, cols=len(table.columns))
        for row in table.rows:
            target_row = target_table.add_row().cells
            for i, cell in enumerate(row.cells):
                target_row[i].text = cell.text

def combine_docx_files(file_paths, output_path):
    # Create a new Document object
    combined_document = Document()

    for i, file_path in enumerate(file_paths):
        doc = Document(file_path)
        if i > 0:
            combined_document.add_page_break()
        append_paragraphs(combined_document, doc)
        append_tables(combined_document, doc)
    
    # Save the combined document
    combined_document.save(output_path)
    print(f"Combined document saved as {output_path}")

def main():
    try:
        file_paths = select_files()
        if file_paths:
            output_file = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                title="Save combined document as")
            if output_file:
                combine_docx_files(file_paths, output_file)
        else:
            print("No files selected")
    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()
    input("Press Enter to exit...")  # Prevent the console from closing immediately
